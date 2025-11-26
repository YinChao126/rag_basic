"""
通用文本提取器模块

该模块提供了从多种文档格式中提取纯文本内容的功能，包括：
- PDF文件
- Excel文件
- Word文档
- Markdown文件
- CSV文件

所有函数都具有统一的接口和错误处理机制。
"""

import os
import pdfplumber
import pandas as pd
import re
import chardet
from docx import Document   #! 此处安装特别注意：uv add python-docx

# API说明：

# 用户唯一需要调用的API是 extract_file，该函数自动根据后缀调用不同的函数，实现文件的提取


def extract_file(file_path: str) -> str:
    """
    通用文本提取入口函数，根据文件后缀自动选择提取函数。
    支持: .pdf, .docx, .md, .csv, .xlsx
    """
    ext = os.path.splitext(file_path)[-1].lower()
    if ext == ".pdf":
        return extract_pdf(file_path)
    elif ext == ".docx":
        return extract_word(file_path)
    elif ext in [".xls", ".xlsx"]:
        return extract_excel(file_path)
    elif ext == ".md":
        return extract_markdown(file_path)
    elif ext == ".csv":
        return extract_csv(file_path)
    else:
        raise ValueError(f"不支持的文件类型: {ext}")
    

def extract_pdf(pdf_path: str) -> str:
    """
    从 PDF 文件中提取文本内容。
    
    使用 pdfplumber 库提取 PDF 文件中的文本内容。

    Args:
        pdf_path (str): PDF 文件的路径

    Returns:
        str: 提取的文本内容，如果出错则返回错误信息
    """

    text_content = ""
    try:
        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"PDF 文件不存在: {pdf_path}")

        print(f"[INFO] 正在尝试从 '{pdf_path}' 提取文本...")
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    text_content += page_text + "\n"
        print(f"[INFO] 成功从 PDF 提取文本。")
    except FileNotFoundError as fnf_err:
        error_msg = f"[ERROR] 文件未找到: {fnf_err}"
        print(error_msg)
        text_content = error_msg
    except Exception as e:
        error_msg = f"[ERROR] 从 PDF '{pdf_path}' 提取文本时出错: {e}"
        print(error_msg)
        import traceback
        traceback.print_exc()
        text_content = error_msg
    return text_content.strip()


def extract_excel(excel_path: str) -> str:
    """
    从 Excel 文件中提取文本内容。
    
    使用 pandas 库提取 Excel 文件中的文本内容。

    Args:
        excel_path (str): Excel 文件的路径

    Returns:
        str: 提取的文本内容，如果出错则返回错误信息
    """
    
    text_content = ""
    try:
        if not os.path.exists(excel_path):
            raise FileNotFoundError(f"Excel 文件不存在: {excel_path}")
            
        print(f"[INFO] 正在尝试从 '{excel_path}' 提取文本...")
        # 读取Excel文件的所有工作表
        excel_data = pd.read_excel(excel_path, sheet_name=None, header=None)
        
        sheet_count = 0
        for sheet_name, df in excel_data.items():
            # 检查工作表是否为空
            if df.empty:
                continue
                
            text_content += f"\n--- Sheet: {sheet_name} ---\n"
            # 将DataFrame转换为字符串
            text_content += df.to_string(index=False, header=False) + "\n"
            sheet_count += 1
            
        if sheet_count == 0:
            text_content = "[INFO] Excel文件中没有数据内容"
        else:
            print(f"[INFO] 成功从 Excel 提取文本。")
    except FileNotFoundError as fnf_err:
        error_msg = f"[ERROR] 文件未找到: {fnf_err}"
        print(error_msg)
        text_content = error_msg
    except Exception as e:
        error_msg = f"[ERROR] 从 Excel '{excel_path}' 提取文本时出错: {e}"
        print(error_msg)
        import traceback
        traceback.print_exc()
        text_content = error_msg
        
    return text_content.strip()


def extract_markdown(md_path: str) -> str:
    """
    从 Markdown 文件中提取纯文本内容。
    
    会移除常见的 Markdown 标记符号以获取更纯净的文本。

    Args:
        md_path (str): Markdown 文件的路径

    Returns:
        str: 提取的纯文本内容，如果出错则返回错误信息
    """
    
    text_content = ""
    try:
        if not os.path.exists(md_path):
            raise FileNotFoundError(f"Markdown 文件不存在: {md_path}")
            
        print(f"[INFO] 正在尝试从 '{md_path}' 提取文本...")
        
        with open(md_path, 'r', encoding='utf-8') as file:
            text_content = file.read()
            
        # 移除Markdown标记以获取纯文本
        # 移除图片标记 [alt](url)
        text_content = re.sub(r'!\[.*?\]\(.*?\)', '', text_content)
        # 移除外链标记 [text](url)
        text_content = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text_content)
        # 移除行内代码 `
        text_content = re.sub(r'`([^`]+)`', r'\1', text_content)
        # 移除代码块标记 ```
        text_content = re.sub(r'```.*?```', '', text_content, flags=re.DOTALL)
        # 移除代码块标记 ```language
        text_content = re.sub(r'```[a-z]*\n.*?```', '', text_content, flags=re.DOTALL)
        # 移除加粗标记 **
        text_content = re.sub(r'\*\*(.*?)\*\*', r'\1', text_content)
        # 移除斜体标记 *
        text_content = re.sub(r'\*([^*]+)\*', r'\1', text_content)
        # 移除标题标记 #
        text_content = re.sub(r'^#+\s*', '', text_content, flags=re.MULTILINE)
        # 移除分隔线 ---
        text_content = re.sub(r'^---$', '', text_content, flags=re.MULTILINE)
        # 移除引用标记 >
        text_content = re.sub(r'^>\s*', '', text_content, flags=re.MULTILINE)
        # 移除列表标记
        text_content = re.sub(r'^[\-*]\s+', '', text_content, flags=re.MULTILINE)
        text_content = re.sub(r'^\d+\.\s+', '', text_content, flags=re.MULTILINE)
        
        print(f"[INFO] 成功从 Markdown 提取文本。")
    except FileNotFoundError as fnf_err:
        error_msg = f"[ERROR] 文件未找到: {fnf_err}"
        print(error_msg)
        text_content = error_msg
    except Exception as e:
        error_msg = f"[ERROR] 从 Markdown '{md_path}' 提取文本时出错: {e}"
        print(error_msg)
        import traceback
        traceback.print_exc()
        text_content = error_msg
        
    return text_content.strip()


def extract_word(word_path: str) -> str:
    """
    从 Word (.docx) 文件中提取文本内容。
    
    使用 python-docx 库提取 Word 文档中的文本内容。

    Args:
        word_path (str): Word 文件的路径

    Returns:
        str: 提取的文本内容，如果出错则返回错误信息
    """
    
    text_content = ""
    try:
        if not os.path.exists(word_path):
            raise FileNotFoundError(f"Word 文件不存在: {word_path}")
            
        print(f"[INFO] 正在尝试从 '{word_path}' 提取文本...")
        
        # 使用 python-docx 读取 .docx 文件
        doc = Document(word_path)
        
        # 提取所有段落的文本
        for paragraph in doc.paragraphs:
            # 只添加非空段落
            if paragraph.text.strip():
                text_content += paragraph.text + "\n"
            
        # 提取表格中的文本
        for table in doc.tables:
            text_content += "\n[表格内容开始]\n"
            for row in table.rows:
                row_content = ""
                for cell in row.cells:
                    row_content += cell.text.strip() + "\t"
                text_content += row_content.rstrip("\t") + "\n"
            text_content += "[表格内容结束]\n"
                
        print(f"[INFO] 成功从 Word 提取文本。")
    except FileNotFoundError as fnf_err:
        error_msg = f"[ERROR] 文件未找到: {fnf_err}"
        print(error_msg)
        text_content = error_msg
    except Exception as e:
        error_msg = f"[ERROR] 从 Word '{word_path}' 提取文本时出错: {e}"
        print(error_msg)
        import traceback
        traceback.print_exc()
        text_content = error_msg
        
    return text_content.strip()


def extract_csv(csv_path: str) -> str:
    """
    从 CSV 文件中提取文本内容。
    
    使用 pandas 库提取 CSV 文件中的文本内容。

    Args:
        csv_path (str): CSV 文件的路径

    Returns:
        str: 提取的文本内容，如果出错则返回错误信息
    """
    
    text_content = ""
    try:
        if not os.path.exists(csv_path):
            raise FileNotFoundError(f"CSV 文件不存在: {csv_path}")
            
        print(f"[INFO] 正在尝试从 '{csv_path}' 提取文本...")
        
        # 尝试不同的编码方式读取 CSV 文件
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
        df = None
        
        for encoding in encodings:
            try:
                # 添加错误处理参数以应对格式问题
                df = pd.read_csv(csv_path, encoding=encoding, on_bad_lines='skip')
                break
            except UnicodeDecodeError:
                continue
            except pd.errors.ParserError:
                continue
        
        if df is None:
            # 如果常规方法失败，则使用二进制模式读取文件并检测编码
            with open(csv_path, 'rb') as file:
                raw_data = file.read()
                encoding = chardet.detect(raw_data)['encoding']
                # 再次尝试读取，使用检测到的编码
                df = pd.read_csv(csv_path, encoding=encoding, on_bad_lines='skip')
        
        # 将DataFrame转换为字符串
        text_content = df.to_string(index=False)
                
        print(f"[INFO] 成功从 CSV 提取文本。")
    except FileNotFoundError as fnf_err:
        error_msg = f"[ERROR] 文件未找到: {fnf_err}"
        print(error_msg)
        text_content = error_msg
    except Exception as e:
        error_msg = f"[ERROR] 从 CSV '{csv_path}' 提取文本时出错: {e}"
        print(error_msg)
        import traceback
        traceback.print_exc()
        text_content = error_msg
        
    return text_content.strip()


if __name__ == "__main__":
    import os
    from pathlib import Path
    
    # 创建输出目录
    output_dir = Path("output")
    output_dir.mkdir(exist_ok=True)
    
    print("=" * 70)
    print("📄 文档提取测试")
    print("=" * 70)
    
    # 测试PDF提取
    pdf_path = "./knowledge/test_knowledge/test_user_manual.pdf"
    
    if not os.path.exists(pdf_path):
        print(f"❌ 错误：文件不存在 {pdf_path}")
        print("💡 提示：请确保 knowledge/test_knowledge/ 目录下有PDF文件")
        exit(1)
    
    print(f"\n📖 正在从PDF提取文本...")
    print(f"   文件路径: {pdf_path}")
    
    pdf_content = extract_pdf(pdf_path)
    
    # 保存到output目录
    output_file = output_dir / "pdf_content.txt"
    with open(output_file, "w", encoding="utf-8") as f:
        f.write(pdf_content)
    
    print(f"\n✅ 提取完成！")
    print(f"   - 提取文本长度: {len(pdf_content)} 字符")
    print(f"   - 输出文件: {output_file}")
    print(f"\n📝 文本预览（前200字符）:")
    print("-" * 70)
    print(pdf_content[:200] + "..." if len(pdf_content) > 200 else pdf_content)
    print("-" * 70)
    


    # # 假设你的 CSV 文件路径为 'data.csv'
    # csv_path = 'knowledge/data.csv'
    # text_content = extract_csv(csv_path)
    # print(text_content)
    
    # # 假设你的 Excel 文件路径为 'data.xlsx'
    # excel_path = 'data.xlsx'
    # text_content = extract_excel(excel_path)